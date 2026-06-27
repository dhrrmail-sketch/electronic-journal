from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "outputs/electronic-journal-tests-beta/Шаблон-полного-курса.docx"
NAVY = RGBColor(21, 43, 77)
BLUE = RGBColor(61, 114, 244)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ШАБЛОН ПОЛНОГО УЧЕБНОГО КУРСА")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY
doc.add_paragraph(
    "Передайте этот файл ИИ вместе с силлабусом и лекциями. "
    "Попросите заполнить поля, не меняя английские служебные метки, и вернуть DOCX."
)

def field(marker, value="", note=""):
    p = doc.add_paragraph()
    r = p.add_run(marker)
    r.bold = True
    r.font.color.rgb = NAVY
    p.add_run(value)
    if note:
        r = p.add_run(f"  [{note}]")
        r.italic = True

def lesson(number):
    p = doc.add_paragraph()
    r = p.add_run(f"LESSON {number}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = BLUE
    field("LECTURE_TOPIC: ", f"Тема лекции {number}")
    field("SEMINAR_TOPIC: ", f"Тема семинара {number}")
    field("SRS_TOPICS: ", f"Тема СРС {number}; дополнительная тема СРС", "несколько тем разделяйте точкой с запятой")
    field("SRSP_TOPICS: ", f"Тема СРСП {number}; дополнительная тема СРСП", "несколько тем разделяйте точкой с запятой")
    field("ESSAY_TOPICS: ", "Тема эссе 1; тема эссе 2", "оставьте пустым, если эссе не требуется")
    field("PRESENTATION_TOPICS: ", "Тема презентации 1; тема презентации 2", "оставьте пустым, если презентация не требуется")
    field("MINDMAP_TOPICS: ", "Тема ментальной карты 1; тема 2", "оценивается; пусто — если не нужно")
    field("INFOGRAPHIC_TOPICS: ", "Тема инфографики 1; тема 2", "оценивается; пусто — если не нужно")
    field("TABLE_TOPICS: ", "Тема таблицы/сравнения 1; тема 2", "оценивается; пусто — если не нужно")
    field("CONSPECTUS_TOPICS: ", "Тема конспекта 1; тема 2", "оценивается; пусто — если не нужно")
    field("OTHER_TASK_TOPICS: ", "Тема реферата; тема творческого задания", "любые дополнительные виды работ")
    field("DEADLINE_DAYS: ", "7", "дней после занятия")
    field("LESSON_PLAN: ", "Вступление; основная часть; закрепление; итоги")
    field("SEMINAR_QUESTIONS: ", "1) Вопрос; 2) Вопрос; 3) Вопрос")
    field(
        "LESSON_PARAGRAPHS: ",
        "Параграф 1 — факты ||| Параграф 2 — понимание ||| Параграф 3 — ценность ||| Параграф 4 — анализ ||| Параграф 5 — творческое осмысление ||| Параграф 6 — вывод",
        "6 параграфов семинара по числу секторов; разделяйте параграфы тремя вертикальными чертами",
    )
    field("TEACHER_NOTES: ", "Методические заметки преподавателю")
    field("RUBRIC: ", "Содержание — 40; аргументация — 30; источники — 20; оформление — 10")
    doc.add_paragraph("END_LESSON")

field("COURSE_TITLE: ", "Название дисциплины")
field("COURSE_CODE: ", "ISL-101")
field("SEMESTER: ", "1")
field("LANGUAGE: ", "ru")
field("DESCRIPTION: ", "Краткое описание курса")
field("LEARNING_OUTCOMES: ", "Результат 1; Результат 2; Результат 3")
for i in range(1, 16):
    lesson(i)

doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("ИНСТРУКЦИЯ ДЛЯ ИИ")
r.bold = True
r.font.size = Pt(17)
r.font.color.rgb = NAVY
for text in [
    "Изучи силлабус и загруженные лекции.",
    "Выстрой 15 лекций и семинаров в логической последовательности.",
    "Для каждого занятия предложи связанные СРС и СРСП.",
    "Добавь эссе, презентации, ментальные карты, инфографику, таблицы и конспекты только там, где они действительно уместны (это СРСП — оцениваемые задания).",
    "Составь план урока, вопросы, рубрику и заметки преподавателю.",
    "Не изменяй служебные метки и строки END_LESSON.",
    "Верни заполненный документ в формате DOCX.",
]:
    doc.add_paragraph(text, style="List Number")

doc.core_properties.title = "Шаблон полного учебного курса"
doc.core_properties.author = "Е. Абдыкаимов"
doc.save(OUT)
print(OUT)
