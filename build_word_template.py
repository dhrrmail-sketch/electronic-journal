from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "outputs/electronic-journal-tests-beta/Шаблон-теста-для-журнала.docx"
NAVY = RGBColor(21, 43, 77)
BLUE = RGBColor(61, 114, 244)
MUTED = RGBColor(95, 107, 124)
PALE = "EEF3FF"


def font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_field(doc, marker, example="", note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(marker)
    font(r, bold=True, color=NAVY)
    if example:
        r = p.add_run(example)
        font(r)
    if note:
        r = p.add_run(f"  [{note}]")
        font(r, size=9, color=MUTED)


def add_question(doc, number, qtype, text, options, correct, points, explanation):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    shade(p, PALE)
    r = p.add_run(f"QUESTION {number}")
    font(r, size=13, bold=True, color=BLUE)
    add_field(doc, "TYPE: ", qtype, "SINGLE, MULTIPLE или TRUE_FALSE")
    add_field(doc, "TEXT: ", text)
    for key, value in options:
        add_field(doc, f"{key}: ", value)
    add_field(doc, "CORRECT: ", correct, "для нескольких ответов: A,C")
    add_field(doc, "POINTS: ", str(points))
    add_field(doc, "EXPLANATION: ", explanation)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("END_QUESTION")
    font(r, bold=True, color=NAVY)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("Электронный журнал • шаблон импорта теста")
font(r, size=9, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Не меняйте служебные метки: TEST_TITLE, QUESTION, TYPE, TEXT, CORRECT, END_QUESTION")
font(r, size=8, color=MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("ШАБЛОН ТЕСТА")
font(r, size=24, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run("Заполните поля и загрузите этот Word-файл в модуль тестирования")
font(r, size=11, color=MUTED)

p = doc.add_paragraph()
shade(p, "FFF4DA")
r = p.add_run("ВАЖНО: ")
font(r, bold=True, color=NAVY)
r = p.add_run("пишите значение после двоеточия. Не удаляйте английские служебные метки и строки END_QUESTION.")
font(r)

add_field(doc, "TEST_TITLE: ", "Тест №1")
add_field(doc, "TEST_TOPIC: ", "Основы вероубеждения")
add_field(doc, "INSTRUCTIONS: ", "Выберите правильный ответ. После завершения нажмите «Отправить тест».")

add_question(
    doc, 1, "SINGLE",
    "Что означает термин «акида»?",
    [("A", "Нравственность"), ("B", "Вероубеждение"), ("C", "История"), ("D", "Право")],
    "B", 1, "Акида — система вероубеждений."
)
add_question(
    doc, 2, "MULTIPLE",
    "Выберите два верных утверждения.",
    [("A", "Первое верное утверждение"), ("B", "Неверное утверждение"),
     ("C", "Второе верное утверждение"), ("D", "Неверное утверждение")],
    "A,C", 2, "Засчитывается только полностью правильный набор."
)
add_question(
    doc, 3, "TRUE_FALSE",
    "Утверждение для проверки является верным.",
    [("A", "Верно"), ("B", "Неверно")],
    "A", 1, "Краткое пояснение правильного ответа."
)

doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("КРАТКАЯ ПАМЯТКА")
font(r, size=18, bold=True, color=NAVY)

steps = [
    "Скопируйте блок QUESTION для каждого нового вопроса.",
    "TYPE: SINGLE — один правильный ответ.",
    "TYPE: MULTIPLE — несколько правильных ответов; в CORRECT укажите, например, A,C.",
    "TYPE: TRUE_FALSE — варианты A: Верно и B: Неверно.",
    "POINTS — количество баллов за вопрос.",
    "EXPLANATION — пояснение, которое студент увидит после разрешённого преподавателем момента.",
    "Сохраните файл как .docx и загрузите его в кабинет тестов.",
]
for i, text in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{i}. ")
    font(r, bold=True, color=BLUE)
    r = p.add_run(text)
    font(r)

doc.core_properties.title = "Шаблон теста для электронного журнала"
doc.core_properties.author = "Е. Абдыкаимов"
doc.core_properties.subject = "Импорт тестов"
doc.save(OUT)
print(OUT)
